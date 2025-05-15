import os
import logging
import docx
import requests
from io import BytesIO
from config import DOCUMENT_PATH

logger = logging.getLogger(__name__)

# Global variables to store document content
document_content = []
document_sections = {}

def extract_text_from_docx(docx_path_or_bytes):
    """Extract text from a .docx file or BytesIO object."""
    try:
        # Check if we received a BytesIO object directly (for Vercel environment)
        if isinstance(docx_path_or_bytes, BytesIO):
            doc = docx.Document(docx_path_or_bytes)
        else:
            # Traditional file path approach
            doc = docx.Document(docx_path_or_bytes)
            
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())
        
        return text
    except Exception as e:
        logger.error(f"Error extracting text from document: {e}")
        return []

def parse_document_structure(content):
    """Parse the document to identify chapters and sections."""
    sections = {}
    current_chapter = None
    current_section = None
    
    for line in content:
        # Check for chapter headings
        if line.startswith("Chapter "):
            parts = line.split(".")
            if len(parts) >= 2:
                current_chapter = parts[1].strip()
                current_section = None
                sections[current_chapter] = {"sections": {}, "content": []}
        # Check for section headings (non-chapter lines that aren't too long)
        elif current_chapter and len(line) < 100 and not line.startswith("Table of Contents") and not line.isdigit():
            current_section = line
            if current_section not in sections[current_chapter]["sections"]:
                sections[current_chapter]["sections"][current_section] = []
        # Add content to current section or chapter
        elif current_chapter:
            if current_section:
                sections[current_chapter]["sections"][current_section].append(line)
            else:
                sections[current_chapter]["content"].append(line)
    
    return sections

def try_vercel_document_loading():
    """Try to load document in a Vercel-compatible way via various methods."""
    try:
        # Check if we're running on Vercel
        is_vercel = os.environ.get('VERCEL') == '1'
        
        if is_vercel or 'VERCEL_URL' in os.environ:
            logger.info("Detected Vercel environment, attempting to load document using various methods")
            
            # Method 1: Try to load from the local filesystem first (our build script should have placed it there)
            vercel_paths = [
                # Standard path for our document
                DOCUMENT_PATH,
                
                # Paths that might be created by our build script
                "/tmp/attached_assets/pharmacy_guide.docx",
                ".vercel/output/static/attached_assets/pharmacy_guide.docx",
                "attached_assets/pharmacy_guide.docx"
            ]
            
            for path in vercel_paths:
                if os.path.exists(path):
                    logger.info(f"Found document at local path: {path}")
                    try:
                        # Attempt to open the file to verify it's readable
                        with open(path, 'rb') as f:
                            sample = f.read(100)  # Read a small sample
                            if sample:  # If we got data, it's likely valid
                                logger.info("Document appears to be valid and readable")
                                return path  # Return the path to the document
                    except Exception as e:
                        logger.error(f"Error reading document at {path}: {e}")
                        # Continue to next path
            
            logger.info("Document not found locally, trying URL-based methods")
            
            # Method 2: Try URL-based loading
            # Get the URL from Vercel environment
            vercel_url = os.environ.get('VERCEL_URL', '')
            if not vercel_url:
                logger.warning("VERCEL_URL environment variable not set")
                # Try an alternate approach using environment variables
                if 'PROJECT_NAME' in os.environ:
                    project_name = os.environ.get('PROJECT_NAME')
                    vercel_url = f"{project_name}.vercel.app"
                    logger.info(f"Using constructed URL from PROJECT_NAME: {vercel_url}")
                else:
                    logger.warning("Unable to determine Vercel URL or project name")
            
            if vercel_url:
                # Construct the URL to the document
                document_url = f"https://{vercel_url}/attached_assets/pharmacy_guide.docx"
                logger.info(f"Attempting to fetch document from: {document_url}")
                
                # Fetch the document
                try:
                    response = requests.get(document_url, timeout=10)
                    if response.status_code == 200:
                        logger.info("Successfully fetched document from URL")
                        document_bytes = BytesIO(response.content)
                        return document_bytes
                    else:
                        logger.error(f"Failed to fetch document: Status {response.status_code}")
                        
                        # Try alternative URLs
                        alt_urls = [
                            f"https://{vercel_url}/attached_assets/pharmacy_guide.docx",
                            f"https://{vercel_url.replace('https://', '')}/attached_assets/pharmacy_guide.docx",
                            f"https://{vercel_url.replace('http://', '')}/attached_assets/pharmacy_guide.docx"
                        ]
                        
                        for alt_url in alt_urls:
                            if alt_url != document_url:  # Skip if same as original URL
                                logger.info(f"Trying alternative URL: {alt_url}")
                                try:
                                    alt_response = requests.get(alt_url, timeout=10)
                                    if alt_response.status_code == 200:
                                        logger.info(f"Successfully fetched document from alternative URL: {alt_url}")
                                        document_bytes = BytesIO(alt_response.content)
                                        return document_bytes
                                except Exception as e:
                                    logger.error(f"Error fetching from alternative URL {alt_url}: {e}")
                except Exception as e:
                    logger.error(f"Error making request to {document_url}: {e}")
            
            # If we get here, all methods failed
            logger.error("All Vercel document loading methods failed")
            return None
            
        return None
    except Exception as e:
        logger.error(f"Error in Vercel document loading: {e}")
        return None

def initialize_document_processor():
    """Initialize the document processor by loading and parsing the document."""
    global document_content, document_sections
    
    # First try Vercel-specific loading
    document_source = try_vercel_document_loading()
    
    if document_source:
        # Check if we got a path or BytesIO object
        if isinstance(document_source, str):
            # We got a file path
            logger.info(f"Loading document from path provided by Vercel handler: {document_source}")
            document_content = extract_text_from_docx(document_source)
        else:
            # We got a BytesIO object
            logger.info("Loading document from BytesIO object provided by Vercel handler")
            document_content = extract_text_from_docx(document_source)
    else:
        # Traditional path - load from file
        if not os.path.exists(DOCUMENT_PATH):
            logger.error(f"Document not found at {DOCUMENT_PATH}")
            
            # Check if we're on Vercel and need to leave a message for debugging
            if os.environ.get('VERCEL') == '1' or 'VERCEL_URL' in os.environ:
                logger.warning("Running on Vercel without access to document file.")
                logger.warning("Please check that the document is properly deployed.")
            
            return False
        
        logger.info(f"Loading document from {DOCUMENT_PATH}")
        document_content = extract_text_from_docx(DOCUMENT_PATH)
    
    if not document_content:
        logger.error("Failed to extract content from document")
        return False
    
    logger.info(f"Successfully loaded document with {len(document_content)} lines")
    
    # Parse document structure
    document_sections = parse_document_structure(document_content)
    logger.info(f"Parsed document into {len(document_sections)} chapters")
    
    return True

def get_document_content():
    """Get the full document content."""
    return document_content

def get_document_sections():
    """Get the parsed document sections."""
    return document_sections

def get_section_content(chapter, section=None):
    """Get content for a specific chapter and section."""
    if chapter in document_sections:
        if section and section in document_sections[chapter]["sections"]:
            return document_sections[chapter]["sections"][section]
        else:
            return document_sections[chapter]["content"]
    return []

def search_document(query, max_results=5):
    """Simple search function to find relevant sections for a query."""
    results = []
    
    # Convert query to lowercase for case-insensitive search
    query_lower = query.lower()
    
    # Search through chapters and sections
    for chapter_name, chapter_data in document_sections.items():
        # Search in chapter content
        chapter_text = " ".join(chapter_data["content"])
        if query_lower in chapter_text.lower():
            results.append({
                "chapter": chapter_name,
                "section": None,
                "relevance": chapter_text.lower().count(query_lower)
            })
        
        # Search in sections
        for section_name, section_content in chapter_data["sections"].items():
            section_text = " ".join(section_content)
            if query_lower in section_text.lower():
                results.append({
                    "chapter": chapter_name,
                    "section": section_name,
                    "relevance": section_text.lower().count(query_lower)
                })
    
    # Sort by relevance and limit results
    results.sort(key=lambda x: x["relevance"], reverse=True)
    return results[:max_results]
